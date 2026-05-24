from pathlib import Path
from shutil import copy2

from PIL import Image
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_BREAK
from docx.shared import Cm, Pt
from docx.oxml.ns import qn


SRC = Path("E:/App/ЗВІТ_Практика_Пистун_оновлений_2.docx")
OUT = Path("E:/App/ЗВІТ_Практика_Пистун_оновлений_3.docx")
DOWNLOADS = Path("C:/Users/pistu/Downloads")
ASSETS = Path("E:/App/practice_report_assets/windows_install_screens")
ASSETS.mkdir(parents=True, exist_ok=True)


def set_run_font(run, size=14, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_heading(doc, text, center=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, bold=True)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    set_run_font(run, size=12)
    return p


def prepare_image(src, dest, max_side=1600):
    img = Image.open(src).convert("RGB")
    scale = min(max_side / max(img.size), 1)
    if scale < 1:
        img = img.resize((round(img.width * scale), round(img.height * scale)), Image.LANCZOS)
    img.save(dest, quality=92, optimize=True)


def add_picture(doc, img_path, caption, width_cm=14.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run()
    run.add_picture(str(img_path), width=Cm(width_cm))
    add_caption(doc, caption)


doc = Document(SRC)

if any(p.text.strip() == "ДОДАТОК Б" for p in doc.paragraphs):
    raise SystemExit("Додаток Б уже є в документі, повторну вставку пропущено.")

screens = [
    ("6.png", "b1_language.jpg", "Рисунок Б.1 - Вибір мовних параметрів у програмі інсталяції Windows 11"),
    ("5.png", "b2_install_type.jpg", "Рисунок Б.2 - Вибір варіанта інсталяції Windows 11"),
    ("4.png", "b3_product_key.jpg", "Рисунок Б.3 - Введення ключа продукту під час інсталяції Windows 11"),
    ("3.png", "b4_disk_location.jpg", "Рисунок Б.4 - Вибір розташування для встановлення Windows 11"),
    ("2.png", "b5_ready.jpg", "Рисунок Б.5 - Перевірка готовності до інсталяції Windows 11 Home"),
    ("1.png", "b6_progress.jpg", "Рисунок Б.6 - Процес встановлення Windows 11"),
]

prepared = []
for src_name, dest_name, caption in screens:
    src = DOWNLOADS / src_name
    if not src.exists():
        raise FileNotFoundError(src)
    dest = ASSETS / dest_name
    prepare_image(src, dest)
    prepared.append((dest, caption))

add_page_break(doc)
add_heading(doc, "ДОДАТОК Б")
add_heading(doc, "Скріншоти етапів встановлення Windows 11", center=False)
add_body(
    doc,
    "У додатку наведено додаткові скріншоти, що деталізують послідовність встановлення Windows 11: вибір мовних параметрів, вибір типу інсталяції, роботу з ключем продукту, вибір дискового простору, підтвердження готовності до інсталяції та початок копіювання системних файлів.",
)

for index, (img_path, caption) in enumerate(prepared):
    if index > 0:
        add_page_break(doc)
    add_picture(doc, img_path, caption)

doc.save(OUT)
print(OUT)
